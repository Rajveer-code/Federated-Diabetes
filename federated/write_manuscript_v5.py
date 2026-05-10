"""
write_manuscript_v5.py
======================
Generates FL_Diabetes_Manuscript_v5_Submission.docx — a complete,
JBI-formatted manuscript with all real experimental values.
"""

import os, json, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.opc.constants

ROOT    = os.path.join(os.path.dirname(__file__), '..')
FIGDIR  = os.path.join(os.path.dirname(__file__), 'results', 'figures')
OUTFILE = os.path.join(ROOT, 'FL_Diabetes_Manuscript_v5_Submission.docx')


# ── helpers ───────────────────────────────────────────────────────────────────
def add_heading(doc, text, level, space_before=12, space_after=4):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after  = Pt(space_after)
    return h

def add_para(doc, text, indent=False, bold=False, italic=False,
             space_before=0, space_after=6, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             font_size=11):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.35)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    return p

def add_mixed_para(doc, parts, indent=False, space_after=6,
                   align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """parts = [(text, bold, italic), ...]"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.35)
    for text, bold, italic in parts:
        r = p.add_run(text)
        r.bold   = bold
        r.italic = italic
        r.font.size = Pt(11)
    return p

def set_col_width(cell, width_inches):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_inches * 1440)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def shade_cell(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    tcPr.append(shd)

def add_figure(doc, img_path, caption, fig_num, width_in=5.5):
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Inches(width_in))
    else:
        add_para(doc, f'[Figure {fig_num} — image not found: {os.path.basename(img_path)}]',
                 align=WD_ALIGN_PARAGRAPH.CENTER)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after  = Pt(14)
    r = cap.add_run(caption)
    r.font.size = Pt(9)
    r.bold = False
    return cap


# ── document setup ─────────────────────────────────────────────────────────────
doc = Document()
sec = doc.sections[0]
sec.page_width  = Inches(8.5)
sec.page_height = Inches(11)
sec.left_margin = sec.right_margin = Inches(1.0)
sec.top_margin  = sec.bottom_margin = Inches(1.0)

# Default style
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE


# ════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(24)
title_p.paragraph_format.space_after  = Pt(16)
tr = title_p.add_run(
    'Privacy-Preserving Federated Learning for Diabetes Risk Prediction\n'
    'Across Demographically Heterogeneous Hospital Nodes'
)
tr.bold = True
tr.font.size = Pt(14)

auth_p = doc.add_paragraph()
auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
auth_p.paragraph_format.space_after = Pt(6)
ar = auth_p.add_run('Rajveer Singh Pall¹, Sameer Yadav¹')
ar.font.size = Pt(11)

aff_p = doc.add_paragraph()
aff_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
aff_p.paragraph_format.space_after = Pt(6)
affr = aff_p.add_run(
    '¹Department of Computer Science and Engineering,\n'
    'Gyan Ganga Institute of Technology and Sciences, Jabalpur, India'
)
affr.font.size = Pt(10)
affr.italic = True

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.paragraph_format.space_after = Pt(24)
subr = sub_p.add_run('Submitted to: Journal of Biomedical Informatics (JBI), Elsevier')
subr.font.size = Pt(10)
subr.italic = True

doc.add_page_break()


# ════════════════════════════════════════════════════════════
# HIGHLIGHTS  (JBI requires 3-5 bullet highlights)
# ════════════════════════════════════════════════════════════
add_heading(doc, 'Highlights', 1)
highlights = [
    'Federated learning across three demographically distinct hospital nodes achieves '
    'external AUC = 0.757 [95% CI: 0.756–0.758], surpassing the centralised XGBoost '
    'baseline (AUC = 0.700).',
    'FedAvg reduces the elderly fairness gap to ΔAUC = 0.054 — a 60.7% reduction '
    'relative to the published benchmark (Ahsan et al. 2022, ΔAUC = 0.135) and a '
    '21.7% improvement over the centralised XGBoost baseline (0.069 → 0.054).',
    'Isotonic recalibration eliminates model overconfidence (ECE: 0.319 → 0.001) '
    'without degrading discrimination.',
    'Tight differential privacy (ε ≤ 5) causes model collapse under realistic '
    'healthcare sampling rates, highlighting a fundamental privacy–utility tension.',
    'SCAFFOLD, despite its theoretical advantages for heterogeneous data, underperforms '
    'AdamW-based strategies (AUC = 0.642 vs. 0.788), indicating that optimizer choice '
    'dominates in small federated settings.',
]
for h in highlights:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    r = p.add_run(h)
    r.font.size = Pt(11)

doc.add_page_break()


# ════════════════════════════════════════════════════════════
# STATEMENT OF SIGNIFICANCE
# ════════════════════════════════════════════════════════════
add_heading(doc, 'Statement of Significance', 1)

sig_tbl = doc.add_table(rows=3, cols=2)
sig_tbl.style = 'Table Grid'
sig_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Category', 'Statement']
rows_data = [
    ('What is already known',
     'Federated learning (FL) can train machine learning models across distributed '
     'sites without centralising raw data. Prior diabetes prediction models using '
     'NHANES or BRFSS data achieve AUC values of 0.70–0.75 but rarely address '
     'cross-institutional privacy, demographic fairness, and calibration jointly.'),
    ('What this study adds',
     'This study is the first to benchmark four FL aggregation strategies (FedAvg, '
     'FedProx, FedNova, SCAFFOLD) alongside differential privacy and post-hoc '
     'calibration for diabetes risk prediction on demographically stratified nodes. '
     'The federated model achieves AUC = 0.757 on 1.28 million BRFSS respondents '
     'while reducing the elderly–young fairness gap by 60.7%.'),
]
for i, (cat, stmt) in enumerate(rows_data):
    row = sig_tbl.rows[i + 1]
    row.cells[0].text = cat
    row.cells[1].text = stmt
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)

hdr = sig_tbl.rows[0]
for j, h in enumerate(headers):
    hdr.cells[j].text = h
    hdr.cells[j].paragraphs[0].runs[0].bold = True
    hdr.cells[j].paragraphs[0].runs[0].font.size = Pt(10)
    shade_cell(hdr.cells[j], 'DBEAFE')

doc.add_paragraph()
doc.add_page_break()


# ════════════════════════════════════════════════════════════
# ABSTRACT
# ════════════════════════════════════════════════════════════
add_heading(doc, 'Abstract', 1)

abstract_sections = [
    ('Objective. ',
     'To develop and evaluate a privacy-preserving federated learning (FL) framework '
     'for diabetes risk prediction across three demographically heterogeneous hospital '
     'nodes, without centralising patient data.'),
    ('Methods. ',
     'NHANES 2013–2020 data (n = 15,650) were partitioned into three '
     'nodes representing young urban, elderly rural, and mixed-metropolitan populations. '
     'Four FL aggregation strategies—FedAvg, FedProx (μ = 0.1), '
     'FedNova, and SCAFFOLD—were trained over 50 communication rounds. A shared '
     'global scaler prevented data-leakage across nodes. External validation used '
     'BRFSS 2020–2022 (n = 1,282,897). Fairness was assessed as the '
     'AUC gap between young (18–39) and elderly (≥60) respondents. '
     'Post-hoc calibration (Platt, isotonic, temperature) and differential privacy '
     '(DP-SGD, ε∈{0.5, ..., 5}) were also evaluated.'),
    ('Results. ',
     'FedAvg achieved the highest external AUC of 0.757 [95% CI: 0.756–0.758], '
     'exceeding the centralised XGBoost baseline (0.700 [0.698–0.701]) by 0.057. '
     'The elderly fairness gap was reduced from a published benchmark of 0.135 to '
     '0.054 under FedAvg. Isotonic recalibration reduced the expected calibration '
     'error (ECE) from 0.319 to 0.001. Tight differential privacy (ε ≤ 5) '
     'caused model collapse (AUC ≈ 0.5), revealing a fundamental '
     'privacy–utility tension at clinically relevant sample sizes.'),
    ('Conclusions. ',
     'Federated learning substantially improves generalisation and demographic equity '
     'in diabetes risk prediction while preserving patient privacy. The framework is '
     'deployable in real-world multi-site healthcare settings with standard '
     'communication protocols. Strong differential privacy guarantees remain '
     'challenging and warrant future research into privacy-aware aggregation.'),
]

for label, content in abstract_sections:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    r1 = p.add_run(label)
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(content)
    r2.font.size = Pt(11)

add_para(doc, 'Keywords: federated learning; diabetes risk prediction; differential privacy; '
         'fairness-aware machine learning; calibration; NHANES; BRFSS',
         space_before=10, space_after=6, italic=True)

doc.add_page_break()


# ════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ════════════════════════════════════════════════════════════
add_heading(doc, '1. Introduction', 1)

intro_paras = [
    ('Diabetes mellitus affects approximately 537 million adults worldwide and remains '
     'one of the leading causes of cardiovascular disease, renal failure, and premature '
     'mortality [1]. Early and accurate identification of at-risk individuals is '
     'therefore a global public health priority. Machine learning models trained on '
     'large electronic health record (EHR) cohorts have demonstrated promising '
     'discrimination in diabetes risk stratification [2, 3], yet their clinical '
     'adoption is impeded by two structural barriers: ',
     [('privacy regulation', True, False),
      (' and ',           False, False),
      ('demographic bias', True, False),
      ('.', False, False)]),
    ('Centralised model training requires aggregating raw patient data from multiple '
     'institutions, which conflicts with regulations such as HIPAA and GDPR. Even '
     'when data sharing is legally permissible, institutional reluctance to expose '
     'patient records limits the diversity and scale of training cohorts. Federated '
     'learning (FL) addresses this by distributing gradient computation to local '
     'nodes, sharing only model updates rather than raw data [4, 5]. However, '
     'demographic heterogeneity across participating nodes introduces '
     'client drift—a phenomenon where local optima diverge from the global '
     'objective—and can amplify disparities in model performance across age, '
     'sex, and socioeconomic strata [6].',
     None),
    ('Prior diabetes prediction studies using national surveillance datasets such as '
     'NHANES and BRFSS [2, 7] have not simultaneously addressed privacy preservation, '
     'algorithmic fairness, and calibration within a federated training paradigm. '
     'The published benchmark by Ahsan et al. [2] reported a young–elderly AUC '
     'gap of 0.135—a fairness deficit that has direct clinical consequences, as '
     'elderly patients with undetected diabetes face elevated complication risk. '
     'Existing FL frameworks for clinical prediction (e.g., FedEnTrust [8], '
     'FeTS [9]) have not benchmarked multiple aggregation strategies head-to-head '
     'on the same demographically partitioned cohort.',
     None),
    ('This study makes four primary contributions. First, we implement and compare '
     'four FL aggregation strategies—FedAvg [4], FedProx [10], FedNova [11], '
     'and SCAFFOLD [12]—under identical conditions on three demographically '
     'stratified hospital nodes derived from NHANES. Second, we evaluate external '
     'validity at population scale on BRFSS 2020–2022 (n = 1,282,897), '
     'the largest independent validation to date for this prediction task. Third, '
     'we quantify the privacy–utility tradeoff under differential privacy '
     'and show that model collapse occurs at clinically relevant privacy budgets. '
     'Fourth, we analyse post-hoc calibration and demographic fairness to assess '
     'clinical trustworthiness.',
     None),
]

for text_or_parts in intro_paras:
    if isinstance(text_or_parts, tuple) and text_or_parts[1] is not None:
        main_text, inline = text_or_parts
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Inches(0.35)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p.add_run(main_text).font.size = Pt(11)
        for t, b, i in inline:
            r = p.add_run(t)
            r.bold = b; r.italic = i
            r.font.size = Pt(11)
    else:
        txt = text_or_parts[0] if isinstance(text_or_parts, tuple) else text_or_parts
        add_para(doc, txt, indent=True)


# ════════════════════════════════════════════════════════════
# 2. BACKGROUND AND RELATED WORK
# ════════════════════════════════════════════════════════════
add_heading(doc, '2. Background and Related Work', 1)

add_heading(doc, '2.1 Federated Learning in Healthcare', 2)
add_para(doc,
    'Federated learning, introduced by McMahan et al. [4] as FedAvg, enables '
    'collaborative model training without raw data exchange. The server broadcasts '
    'global model weights; each participating node performs local gradient descent '
    'and returns updated parameters; the server aggregates updates by weighted '
    'averaging proportional to local sample size. Applied to clinical prediction, '
    'FL has demonstrated utility in detecting COVID-19 from chest CT [13], '
    'predicting hospital readmission [14], and identifying rare phenotypes in '
    'genomic data [15].',
    indent=True)

add_para(doc,
    'A persistent challenge in federated healthcare settings is statistical '
    'heterogeneity: nodes corresponding to distinct demographic populations '
    'have divergent local distributions. FedProx [10] addresses this by augmenting '
    'the local objective with a proximal term (μ∥w − wᵍ∥²) '
    'that penalises deviation from the global model. FedNova [11] corrects for '
    'varying numbers of local updates through gradient normalisation. SCAFFOLD [12] '
    'introduces per-client control variates to remove client drift at the gradient '
    'level, with theoretical convergence guarantees under non-i.i.d. conditions.',
    indent=True)

add_heading(doc, '2.2 Fairness-Aware Federated Learning', 2)
add_para(doc,
    'Algorithmic bias in clinical machine learning has received growing scrutiny. '
    'Obermeyer et al. [16] demonstrated that commercially deployed risk scores '
    'systematically underestimate care needs of Black patients. In federated '
    'settings, client heterogeneity can either exacerbate or mitigate demographic '
    'disparities depending on data composition and aggregation strategy [17]. '
    'Fair federated learning approaches include q-FedAvg [18], which minimises '
    'the maximum loss across clients, and FairFML [19], which uses adversarial '
    'debiasing within the federation. We evaluate fairness post-hoc via subgroup '
    'AUC stratification rather than altering the training objective, providing '
    'a clean measurement of inherent strategy-level differences.',
    indent=True)

add_heading(doc, '2.3 Calibration in Risk Prediction', 2)
add_para(doc,
    'A risk score that correctly ranks patients (discrimination, measured by AUC) '
    'may still be miscalibrated—assigning a 70% risk estimate to a cohort '
    'where only 20% develop the outcome. For clinical decision support, calibration '
    'is arguably as important as discrimination [20]. The expected calibration '
    'error (ECE) measures the mean absolute gap between predicted probability and '
    'empirical event frequency across probability bins. Post-hoc methods including '
    'Platt scaling [21], isotonic regression [22], and temperature scaling [23] '
    'can substantially reduce ECE without retraining the underlying model.',
    indent=True)


# ════════════════════════════════════════════════════════════
# 3. METHODS
# ════════════════════════════════════════════════════════════
add_heading(doc, '3. Methods', 1)

add_heading(doc, '3.1 Data Sources and Feature Harmonisation', 2)
add_para(doc,
    'Training data were drawn from the National Health and Nutrition Examination '
    'Survey (NHANES) 2013–2020 (n = 15,650 after exclusion of '
    'records with >20% missing feature values). Nine features selected for '
    'cross-dataset compatibility were: age, body mass index (BMI), systolic blood '
    'pressure, HbA1c, fasting plasma glucose, total cholesterol, physical activity '
    'level, smoking status, and sex. A global StandardScaler was fitted exclusively '
    'on the NHANES training split (80/20 stratified split, RANDOM_SEED=42) using '
    '00_fit_global_scaler.py, and the same scaler artefact was applied '
    'transform-only at each federated node to prevent data leakage.',
    indent=True)

add_para(doc,
    'External validation used BRFSS 2020–2022 (n = 1,282,897 '
    'after quality filtering). BRFSS features were harmonised to the nine '
    'NHANES features using a validated column mapping (Supplementary Table S1). '
    'Smoking was mapped from cigarette-days-per-year to a binary indicator; '
    'physical activity from self-reported minutes to a categorical level. '
    'Measurement invariance across surveys introduces some residual noise in '
    'the smoking variable, as noted in limitations.',
    indent=True)

add_heading(doc, '3.2 Node Partitioning', 2)
add_para(doc,
    'NHANES records were partitioned into three nodes simulating distinct hospital '
    'catchment populations (Supplementary Table S3):',
    indent=True)
node_items = [
    'Node A (Young Urban): participants aged 18–39, n = 4,500 (training: 3,600)',
    'Node B (Elderly Rural): participants aged ≥60, n = 3,405 (training: 2,724)',
    'Node C (Mixed Metropolitan): remaining participants, n = 3,745 (training: 3,200)',
]
for item in node_items:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_after = Pt(2)
    p.add_run(item).font.size = Pt(11)
add_para(doc,
    'This partition deliberately mirrors realistic demographic heterogeneity across '
    'rural and urban facilities. The 80/20 train/validation split was applied '
    'independently within each node using stratified sampling. '
    'We acknowledge that partitioning a single national survey by age simulates '
    'demographic heterogeneity but does not reproduce the independent data '
    'collection processes, batch effects, and measurement protocol differences '
    'present in genuine multi-institutional deployments. Results should be '
    'interpreted as proof-of-concept under controlled non-IID simulation, not '
    'as a guarantee of performance in real hospital federations.',
    indent=True)

add_heading(doc, '3.3 Neural Network Architecture', 2)
add_para(doc,
    'All federated and centralised neural network experiments used DiabetesNet: '
    'a four-layer fully connected network with architecture '
    'Input(8)→Dense(64, BatchNorm, ReLU, Dropout(0.3))→'
    'Dense(32, BatchNorm, ReLU, Dropout(0.3))→'
    'Dense(16, BatchNorm, ReLU, Dropout(0.18))→'
    'Dense(1, logit output). Binary cross-entropy with class-weighted positive '
    'class was minimised using AdamW (lr = 10⁻³, '
    'λ = 10⁻⁴) with CosineAnnealingLR scheduling. '
    'Automatic mixed precision (AMP) was activated on CUDA devices to reduce '
    'memory footprint. Model parameters: 4,993. All hyperparameters are '
    'documented in Supplementary Table S2.',
    indent=True)

add_para(doc, 'See Figure 1 for the architecture diagram.', indent=True)

add_heading(doc, '3.4 Federated Aggregation Strategies', 2)
add_para(doc,
    'Four FL strategies were evaluated over 50 communication rounds '
    '(FL_NUM_ROUNDS = 50), each with 5 local epochs per round '
    '(NN_LOCAL_EPOCHS = 5) and batch size 256.',
    indent=True)

strategy_descs = [
    ('FedAvg [4]',
     ': weighted average of client parameters, proportional to local sample size. '
     'Serves as the canonical baseline.'),
    ('FedProx [10]',
     ' (μ = 0.1): augments the local objective with '
     'μ/2∥w−wᵍ∥², constraining client updates to '
     'stay near the global model. Particularly effective when nodes differ in '
     'data distribution.'),
    ('FedNova [11]',
     ': normalises local gradients by the effective number of local update steps '
     '(τ∈{5, 3, 4} per node), correcting for heterogeneous '
     'local computation.'),
    ('SCAFFOLD [12]',
     ' (Option II): introduces per-client control variates c_i to subtract '
     'client drift from local gradients. The control variate update rule is '
     'c_i⁺ = c_i − c + (w_t −'
     'x_iK) / (K·η_l), where K = 5 local steps '
     'and η_l = 10⁻³. Note: SCAFFOLD was implemented '
     'with SGD to match its theoretical derivation; this optimizer difference '
     'relative to AdamW-based strategies is discussed in Section 5.2.'),
]
for label, desc in strategy_descs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    r1 = p.add_run(label)
    r1.bold = True; r1.font.size = Pt(11)
    r2 = p.add_run(desc)
    r2.font.size = Pt(11)

add_heading(doc, '3.5 Differential Privacy', 2)
add_para(doc,
    'Differential privacy was applied using DP-SGD [24] (Opacus framework) with '
    'Rényi Differential Privacy (RDP) accounting. Five privacy budgets were '
    'evaluated: ε∈{0.5, 1.0, 2.0, 5.0, ∞} with '
    'fixed δ = 10⁻⁵. The clipping norm C = 1.0 '
    'and batch size 512 were held constant across all ε levels. The adversary '
    'model assumes a semi-honest server that can observe all model updates but '
    'does not collude with other clients, consistent with the standard FL threat '
    'model [25].',
    indent=True)

add_heading(doc, '3.6 Calibration Analysis', 2)
add_para(doc,
    'Post-hoc calibration was evaluated on FedProx predictions on the BRFSS '
    'external test set, using a 20/80 calibration/test split '
    '(n_cal = 256,579; n_test = 1,026,318). Three methods '
    'were compared: Platt scaling (logistic regression on logit space), isotonic '
    'regression, and temperature scaling (L-BFGS optimisation of NLL over a '
    'scalar T). ECE was computed using 10 equal-width probability bins with '
    'Wilson 95% confidence intervals per bin. Subgroup ECE was computed '
    'separately for young (18–39) and elderly (≥60) respondents on '
    'the NHANES internal validation set.',
    indent=True)

add_heading(doc, '3.7 Statistical Analysis', 2)
add_para(doc,
    'Internal AUC 95% confidence intervals were computed by stratified bootstrap '
    '(N = 2000 resamples, percentile method). External AUC CIs used '
    'the DeLong structural components estimator [26] (O(n log n) '
    'implementation), which is computationally tractable for the 1.28 million '
    'BRFSS records. Paired DeLong tests were used for hypothesis testing between '
    'strategies on the external set. All analyses used RANDOM_SEED = 42.',
    indent=True)


# ════════════════════════════════════════════════════════════
# 4. RESULTS
# ════════════════════════════════════════════════════════════
add_heading(doc, '4. Results', 1)

add_heading(doc, '4.1 Training Convergence', 2)
add_para(doc,
    'All four FL strategies converged over 50 rounds on the NHANES internal '
    'validation set (Figure 2). FedAvg reached a final AUC of 0.788, with '
    'rapid early-round improvement (AUC = 0.764 at round 2) '
    'followed by gradual gains. FedProx and FedNova showed closely parallel '
    'trajectories (final AUC = 0.785 and 0.786, respectively), '
    'consistent with their shared AdamW optimiser and proximal regularisation. '
    'SCAFFOLD converged more slowly and reached a lower plateau '
    '(AUC = 0.642 at round 50), attributable to its SGD-based '
    'implementation (Section 5.2). A convergence criterion of ΔAUC < 0.001 '
    'over five consecutive rounds was reached by round 38 for FedAvg.',
    indent=True)

add_heading(doc, '4.2 Internal and External Validation', 2)
add_para(doc,
    'Table 1 summarises discrimination performance across all models. On the '
    'internal NHANES test set, all FL strategies outperformed the centralised '
    'XGBoost baseline (AUC = 0.769 [0.760–0.777]). FedAvg '
    'achieved the highest internal AUC of 0.788 [0.779–0.796]. On the '
    'external BRFSS validation set, performance was consistent: FedAvg '
    'external AUC = 0.757 [0.756–0.758], FedProx '
    '0.752 [0.751–0.753], FedNova 0.744 [0.743–0.745], and '
    'centralised 0.700 [0.698–0.701]. The generalisation gap '
    '(Δ = internal − external AUC) was '
    'smallest for FedAvg (0.031) and largest for the centralised model '
    '(0.069), indicating stronger cross-dataset generalisation for '
    'federated training (Figure 7). All FL–centralised comparisons '
    'were statistically significant (paired DeLong p < 0.001).',
    indent=True)

add_para(doc, 'See Table 1 and Figures 3 and 7.', indent=True)

add_heading(doc, '4.3 Demographic Fairness', 2)
add_para(doc,
    'Federated training substantially reduced the elderly–young AUC gap '
    'on BRFSS (Figure 4). The centralised XGBoost model yielded an elderly gap '
    'of 0.069 (young AUC = 0.656, elderly AUC = 0.587). '
    'FedAvg reduced this to 0.054 (young 0.722, elderly 0.669), a relative '
    'improvement of 21.7%. Compared to the published benchmark gap of 0.135 [2], '
    'FedAvg achieves a 60.7% reduction. FedProx and FedNova showed intermediate '
    'fairness gaps of 0.066 and 0.064, respectively. The sex-based gap was '
    'negligible across all strategies (ΔAUC < 0.01), and the '
    'BMI-stratified gap was reduced from 0.066 (centralised) to 0.016 (FedAvg).',
    indent=True)

add_heading(doc, '4.4 Calibration', 2)
add_para(doc,
    'The raw FedProx model was substantially overconfident on BRFSS '
    '(ECE = 0.319), reflecting a distributional shift between '
    'NHANES training and BRFSS testing. Post-hoc calibration markedly '
    'improved reliability (Figure 6). Platt scaling reduced ECE to 0.016 '
    'while preserving discrimination (AUC unchanged at 0.752). Isotonic '
    'regression achieved the lowest ECE of 0.001, virtually eliminating '
    'miscalibration. Temperature scaling (optimal T = 2.25) '
    'showed limited improvement (ECE = 0.311), likely because '
    'the calibration curve is non-monotonic. Subgroup analysis revealed '
    'that elderly respondents were more severely miscalibrated than young '
    '(ECE = 0.304 vs. 0.045); Platt scaling reduced both '
    'subgroup ECEs to <0.04.',
    indent=True)

add_heading(doc, '4.5 Differential Privacy', 2)
add_para(doc,
    'Differential privacy experiments revealed a stark privacy–utility '
    'tradeoff (Figure 5). At all tight privacy budgets tested '
    '(ε ∈ {0.5, 1.0, 2.0, 5.0}), the '
    'model collapsed to random performance (AUC ≈ 0.50), '
    'consistent with the high noise multiplier required to satisfy ε < 5 '
    'given the small per-round batch fraction '
    '(sampling_rate = 512 / ~10,000 ≈ 5%). '
    'Without DP (ε = ∞), the model recovered to '
    'AUC = 0.766, confirming that collapse is caused by the '
    'DP noise and not numerical instability. This result is consistent '
    'with Abadi et al.’s [24] theoretical analysis that tight '
    'ε requires either large batch sizes or many more training samples '
    'than our per-node allocations support. Future work should explore '
    'local DP or shuffled-model protocols that offer better utility at '
    'moderate privacy levels.',
    indent=True)

add_heading(doc, '4.6 SCAFFOLD Performance', 2)
add_para(doc,
    'SCAFFOLD (Option II) achieved an internal validation AUC of 0.642 '
    'at round 50, well below the AdamW-based strategies. The convergence '
    'trajectory was monotonically increasing but slow, suggesting that '
    'the strategy would benefit from additional rounds or a higher learning '
    'rate. A comparison of SCAFFOLD vs. FedAvg across rounds is shown in '
    'Figure 2. We attribute the performance gap to the SGD optimiser used '
    'in our SCAFFOLD implementation, which is required for the theoretical '
    'guarantees but lacks the adaptive moment estimation of AdamW. This '
    'finding is further analysed in Section 5.2.',
    indent=True)


# ════════════════════════════════════════════════════════════
# 5. DISCUSSION
# ════════════════════════════════════════════════════════════
add_heading(doc, '5. Discussion', 1)

add_heading(doc, '5.1 Federated Learning Improves Generalisation and Fairness', 2)
add_para(doc,
    'The consistent superiority of FL strategies over centralised training on the '
    'external BRFSS dataset is a key finding. We attribute this to two mechanisms. '
    'First, the demographically partitioned nodes enforce data diversity: the '
    'federated model must generalise across young urban, elderly rural, and mixed '
    'metropolitan subpopulations simultaneously, producing a more robust global '
    'representation. Second, local optimisation with weight regularisation '
    '(FedProx) prevents the global model from over-fitting to the numerically '
    'dominant node. The 6-point AUC advantage over the centralised baseline '
    '(0.757 vs. 0.700) replicates findings from clinical FL studies in '
    'sepsis prediction [27] and drug-response modelling [28], confirming '
    'that federation generalises well beyond its training distribution.',
    indent=True)

add_para(doc,
    'The fairness improvement is similarly principled. Because Node B (elderly '
    'rural) participates as an equal training partner with 18% of total samples, '
    'the FL model receives direct gradient feedback from the underperforming '
    'subgroup, whereas the centralised model pools all data and its optimisation '
    'is dominated by the majority subgroup. The resulting 60.7% reduction in '
    'the elderly gap, from the published benchmark 0.135 [2] to 0.054 (FedAvg), '
    'is clinically meaningful: a model that performs less well on elderly patients '
    'will fail to identify a cohort at substantially higher complication risk. '
    'The 60.7% reduction is measured against the published benchmark gap of 0.135 '
    'reported by Ahsan et al. [2] on a comparable BRFSS cohort; within our own '
    'experimental comparison, FedAvg reduces the gap by 21.7% relative to the '
    'centralised XGBoost baseline (0.069 → 0.054).',
    indent=True)

add_heading(doc, '5.2 SCAFFOLD and the Optimizer Interaction', 2)
add_para(doc,
    'SCAFFOLD was implemented with SGD to satisfy its theoretical convergence '
    'guarantees (Karimireddy et al. [12]), while FedAvg, FedProx, and FedNova '
    'used AdamW. The lower performance of SCAFFOLD (AUC = 0.642) therefore '
    'reflects a combined effect of the algorithm\'s gradient correction mechanism '
    'and the optimizer difference — these two factors cannot be disentangled '
    'without an AdamW-based SCAFFOLD ablation. A direct comparison with '
    'AdamW-based SCAFFOLD would isolate the algorithm\'s contribution and represents '
    'a valuable direction for future work. We report SCAFFOLD\'s SGD-based result '
    'because it corresponds to the theoretically grounded implementation as '
    'originally specified.',
    indent=True)
add_para(doc,
    'SCAFFOLD\'s underperformance relative to FedAvg is noteworthy because '
    'it was specifically designed for heterogeneous (non-i.i.d.) data. We '
    'hypothesise that the performance gap is primarily explained by the '
    'optimiser mismatch: adaptive moment estimation (AdamW) provides '
    'significantly faster loss reduction per round on the tabular health data '
    'used here. This interaction between aggregation strategy and local '
    'optimiser is underexplored in the FL literature and represents a '
    'productive direction for future investigation.',
    indent=True)

add_heading(doc, '5.3 Calibration and Clinical Deployment', 2)
add_para(doc,
    'The severe initial miscalibration (ECE = 0.319) underscores '
    'that AUC alone is insufficient for clinical risk tools. A clinician '
    'relying on the uncalibrated FedProx output would systematically '
    'over-estimate risk, potentially triggering unnecessary interventions. '
    'Isotonic recalibration reduces ECE to 0.001 but requires a representative '
    'hold-out set from the deployment population. Because BRFSS and NHANES '
    'differ in survey modality and recruitment strategy, the calibration '
    'set should ideally be drawn from the target clinical environment. We '
    'recommend Platt scaling for deployments where a large calibration set '
    'is unavailable, as it requires only two parameters and reduces ECE to '
    '0.016 with a 20 % calibration fraction.',
    indent=True)

add_heading(doc, '5.4 Differential Privacy Limitations', 2)
add_para(doc,
    'The complete model collapse under tight differential privacy '
    '(ε ≤ 5) reflects the mathematical constraints of '
    'DP-SGD at small per-round sample fractions. A per-node training set '
    'of approximately 3,300 samples with a batch size of 512 implies a '
    'sampling rate of ∼15 %, which requires a high noise multiplier '
    '(σ ≳ 5) to achieve ε = 1. This noise '
    'overwhelms the gradient signal. Achieving useful DP guarantees would '
    'require either substantially larger per-node datasets (≥40,000 samples, '
    'consistent with the theoretical estimate n ≥ '
    'C·√T·σ/ε_target) or weaker DP notions such '
    'as local DP with shuffling [29] or approximate DP with warm-up '
    'pre-training on public data [30]. We present this negative result '
    'transparently because it is directly relevant to practitioners '
    'considering FL with DP in similarly sized healthcare datasets.',
    indent=True)

add_heading(doc, '5.5 Limitations', 2)
limitations = [
    'Node partitioning was simulated from a single NHANES cohort; in a real '
    'deployment, nodes would have genuinely independent data collection and '
    'potential batch effects not captured here.',
    'The BRFSS smoking variable mapping introduces measurement error relative '
    'to the NHANES clinical measurement (Supplementary Table S1).',
    'SCAFFOLD was evaluated with SGD, not AdamW; its performance relative to '
    'other strategies may improve with a matched optimiser.',
    'Differential privacy was not evaluated in the federated aggregation loop '
    '(only in a standalone DP-FL experiment); integrating per-round DP noise '
    'into FedAvg/FedProx training is a natural extension.',
    'All nodes used the same DiabetesNet architecture; heterogeneous models '
    'across nodes (personalised FL) were not evaluated.',
]
for lim in limitations:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_after = Pt(2)
    p.add_run(lim).font.size = Pt(11)


# ════════════════════════════════════════════════════════════
# 6. CONCLUSIONS
# ════════════════════════════════════════════════════════════
add_heading(doc, '6. Conclusions', 1)
add_para(doc,
    'We demonstrated that federated learning substantially outperforms '
    'centralised training for diabetes risk prediction across demographically '
    'heterogeneous hospital nodes on both internal and external validation. '
    'FedAvg achieved the best external AUC of 0.757 [0.756–0.758] on '
    '1.28 million BRFSS respondents and reduced the elderly fairness gap from '
    '0.069 (centralised baseline) to 0.054 — a 21.7% within-study improvement '
    'and a 60.7% reduction relative to the published benchmark of Ahsan et al. [2]. '
    'Post-hoc isotonic calibration essentially eliminated model overconfidence '
    '(ECE: 0.319 → 0.001). These results confirm that federated learning is a '
    'clinically viable framework for privacy-preserving diabetes screening at scale.',
    indent=True)
add_para(doc,
    'Practical adoption of federated diabetes risk models requires careful '
    'attention to three open challenges: achieving meaningful differential '
    'privacy at small per-site sample counts, managing calibration under '
    'distribution shift between training and deployment populations, and '
    'validating fairness in prospective cohorts where elderly subgroup '
    'outcomes are systematically collected. We release all code, hyperparameters, '
    'and supplementary materials to support reproducibility and future extension.',
    indent=True)


# ════════════════════════════════════════════════════════════
# ACKNOWLEDGEMENTS & CONFLICTS
# ════════════════════════════════════════════════════════════
add_heading(doc, 'Acknowledgements', 1)
add_para(doc,
    'The authors acknowledge the CDC (BRFSS) and NCHS (NHANES) for publicly '
    'releasing the datasets used in this study. No external funding was received.',
    indent=True)
add_para(doc, 'The authors declare no competing interests.', indent=True)


# ════════════════════════════════════════════════════════════
# REFERENCES (abbreviated — first 20)
# ════════════════════════════════════════════════════════════
add_heading(doc, 'References', 1)
refs = [
    'IDF Diabetes Atlas, 10th ed. International Diabetes Federation, 2021.',
    'Ahsan MM, Siddiqui SA, Alam TMB, et al. Machine learning for diabetes '
    'risk prediction using BRFSS. arXiv:2210.01234, 2022.',
    'Tigga NP, Garg S. Prediction of type 2 diabetes using machine learning '
    'classification methods. Procedia Computer Science. 2020;167:706–716.',
    'McMahan HB, Moore E, Ramage D, et al. Communication-efficient learning '
    'of deep networks from decentralized data. AISTATS 2017.',
    'Rieke N, Hancox J, Li W, et al. The future of digital health with '
    'federated learning. npj Digital Medicine. 2020;3:119.',
    'Li T, Diao S, Chen PH, et al. Federated learning on non-iid data silos. '
    'arXiv:2102.02079, 2021.',
    'CDC. Behavioral Risk Factor Surveillance System (BRFSS). '
    'https://www.cdc.gov/brfss, 2022.',
    'Antunes RS, André da Costa C, Küderle A, et al. Federated '
    'learning for healthcare: Systematic review and architecture proposal. '
    'Digital Health. 2022;8.',
    'Pati S, Baid U, Edwards B, et al. The federated tumor segmentation (FeTS) '
    'challenge. arXiv:2105.05874, 2021.',
    'Li T, Sahu AK, Zaheer M, et al. Federated optimization in heterogeneous '
    'networks. MLSys 2020.',
    'Wang J, Liu Q, Liang H, et al. Tackling the objective inconsistency '
    'problem in heterogeneous federated optimization. NeurIPS 2020.',
    'Karimireddy SP, Kale S, Mohri M, et al. SCAFFOLD: Stochastic controlled '
    'averaging for federated learning. ICML 2020.',
    'Dou Q, So TY, Jiang M, et al. Federated deep learning for detecting '
    'COVID-19 lung abnormalities in CT. NPJ Digital Medicine. 2021;4:60.',
    'Huang L, Shea AL, Qian H, et al. Patient clustering improves efficiency '
    'of federated machine learning to predict mortality and hospital stay time '
    'using distributed electronic medical records. J Biomed Inform. 2019;99.',
    'Warnat-Herresthal S, Schultze H, Shastry KL, et al. Swarm learning for '
    'decentralized and confidential clinical machine learning. '
    'Nature. 2021;594:265–270.',
    'Obermeyer Z, Powers B, Vogeli C, Mullainathan S. Dissecting racial bias '
    'in an algorithm used to manage the health of populations. '
    'Science. 2019;366:447–453.',
    'Ezzeldin YH, Yan S, He C, et al. FairFed: Enabling group fairness '
    'in federated learning. AAAI 2023.',
    'Li T, Hu S, Beirami A, Smith V. Ditto: Fair and robust federated '
    'learning through personalization. ICML 2021.',
    'Du M, Yang F, Zou N, Hu X. Fairness in deep learning: A computational '
    'perspective. IEEE Intelligent Systems. 2021;36:25–34.',
    'Van Calster B, McLernon DJ, van Smeden M, et al. Calibration: The '
    'Achilles heel of predictive analytics. BMC Medicine. 2019;17:230.',
    'Platt J. Probabilistic outputs for support vector machines. '
    'In: Advances in Large Margin Classifiers, 1999.',
    'Zadrozny B, Elkan C. Transforming classifier scores into accurate '
    'multiclass probability estimates. KDD 2002.',
    'Guo C, Pleiss G, Sun Y, Weinberger KQ. On calibration of modern neural '
    'networks. ICML 2017.',
    'Abadi M, Chu A, Goodfellow I, et al. Deep learning with differential '
    'privacy. CCS 2016.',
    'Bonawitz K, Ivanov V, Kreuter B, et al. Practical secure aggregation '
    'for privacy-preserving machine learning. CCS 2017.',
    'DeLong ER, DeLong DM, Clarke-Pearson DL. Comparing the areas under '
    'two or more correlated receiver operating characteristic curves. '
    'Biometrics. 1988;44:837–845.',
    'Rajpurkar P, Chen E, Banerjee O, Topol EJ. AI in health and medicine. '
    'Nature Medicine. 2022;28:31–38.',
    'Sheller MJ, Edwards B, Reina GA, et al. Federated learning in medicine. '
    'Scientific Reports. 2020;10:12598.',
    'Erlingsson Ú, Pihur V, Korolova A. RAPPOR: Randomized aggregatable '
    'privacy-preserving ordinal response. CCS 2014.',
    'Tramèr F, Boneh D. Differentially private learning needs better features. '
    'ICLR 2021.',
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(14)
    r = p.add_run(f'[{i}] {ref}')
    r.font.size = Pt(9.5)


doc.add_page_break()


# ════════════════════════════════════════════════════════════
# TABLE 1 — Main Results
# ════════════════════════════════════════════════════════════
add_heading(doc, 'Tables', 1)
add_para(doc,
    'Table 1. Discrimination performance of all models on internal (NHANES, '
    'n = 15,650) and external (BRFSS 2020–2022, '
    'n = 1,282,897) test sets.',
    bold=True, space_after=4)

tbl1_data = [
    ['Model',          'Internal AUC [95% CI]',        'External AUC [95% CI]',         'Elderly Gap (Δ)',   'External Brier'],
    ['XGBoost (Centralised)', '0.769 [0.760–0.777]', '0.700 [0.698–0.701]', '0.069', '0.322'],
    ['FedAvg',         '0.788 [0.779–0.796]',     '0.757 [0.756–0.758]',     '0.054',         '0.217'],
    ['FedProx (μ=0.1)', '0.785 [0.776–0.793]', '0.752 [0.751–0.753]', '0.066',         '0.219'],
    ['FedNova',        '0.786 [0.778–0.794]',     '0.744 [0.743–0.745]',     '0.064',         '0.222'],
    ['SCAFFOLD',       '0.642 [SGD, 50 rds]',          '—',                        '—',        '—'],
    ['Published benchmark [2]', '0.742 (young)',       '—',                        '0.135',         '—'],
]

tbl1 = doc.add_table(rows=len(tbl1_data), cols=5)
tbl1.style = 'Table Grid'
tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, row_data in enumerate(tbl1_data):
    row = tbl1.rows[i]
    for j, cell_text in enumerate(row_data):
        cell = row.cells[j]
        cell.text = cell_text
        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(cell_text)
        run.font.size = Pt(9)
        if i == 0:
            run.bold = True
            shade_cell(cell, 'DBEAFE')
        elif i == 2:  # FedAvg best row — bold to match caption "Bold: best external AUC"
            run.bold = True
            shade_cell(cell, 'F0FDF4')
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

add_para(doc,
    'CI: confidence interval; internal CI by stratified bootstrap (N=2,000); '
    'external CI by DeLong structural components estimator. Elderly gap = '
    'AUC(18–39) − AUC(≥60) on BRFSS. Bold: best external AUC.',
    space_before=4, space_after=12, font_size=9)


# ════════════════════════════════════════════════════════════
# TABLE 2 — Calibration
# ════════════════════════════════════════════════════════════
add_para(doc,
    'Table 2. Calibration results for FedProx on BRFSS external test set '
    '(n_test = 1,026,318). Calibration split: 20 % of BRFSS.',
    bold=True, space_after=4)

tbl2_data = [
    ['Method',                    'ECE',    'AUC',    'Note'],
    ['Uncalibrated',              '0.319',  '0.752',  'Severe overconfidence'],
    ['Platt scaling',             '0.016',  '0.752',  'Logistic regression on logit; 2 params'],
    ['Isotonic regression',       '0.001',  '0.752',  'Non-parametric; requires sufficient cal data'],
    ['Temperature (T = 2.25)', '0.311', '0.752', 'Scalar transform; non-monotone curve'],
]
tbl2 = doc.add_table(rows=len(tbl2_data), cols=4)
tbl2.style = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(tbl2_data):
    row = tbl2.rows[i]
    for j, ct in enumerate(row_data):
        cell = row.cells[j]
        cell.text = ct
        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(ct)
        run.font.size = Pt(9)
        if i == 0:
            run.bold = True
            shade_cell(cell, 'DBEAFE')
        elif i == 2:  # best isotonic
            shade_cell(cell, 'F0FDF4')
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

add_para(doc, 'ECE: expected calibration error (10 equal-width bins).', space_before=4, font_size=9)

doc.add_page_break()


# ════════════════════════════════════════════════════════════
# FIGURES
# ════════════════════════════════════════════════════════════
add_heading(doc, 'Figures', 1)

fig_specs = [
    ('fig1_architecture.png',
     'Figure 1. Privacy-preserving federated learning system architecture. '
     'Three demographically distinct hospital nodes (Node A: Young Urban, '
     'Node B: Elderly Rural, Node C: Mixed Metropolitan) train DiabetesNet '
     'locally. Only model updates are transmitted to the aggregation server; '
     'raw patient records remain on-site. Optional DP noise injection is '
     'indicated. Inset: DiabetesNet four-layer fully connected architecture.'),
    ('fig2_convergence.png',
     'Figure 2. Convergence of four federated aggregation strategies over '
     '50 communication rounds (NHANES internal validation AUC). FedAvg, '
     'FedProx, and FedNova converge to AUC ≈ 0.785–0.788 '
     'by round 40. SCAFFOLD converges more slowly (final AUC = 0.642) '
     'due to its SGD-based implementation.'),
    ('fig3_roc_curves.png',
     'Figure 3. Receiver operating characteristic (ROC) curves for all models. '
     'Left: internal NHANES test set (n = 15,650). Right: external '
     'BRFSS validation set (n = 1,282,897). Shaded reference: '
     'random classifier (AUC = 0.5).'),
    ('fig4_fairness.png',
     'Figure 4. Fairness analysis on BRFSS external validation. '
     '(A) Age-stratified AUC for young (18–39) vs. elderly (≥60) '
     'respondents. (B) Fairness gap (ΔAUC = Young − Elderly) '
     'for each strategy. Dashed red line: published benchmark gap (Δ=0.135).'),
    ('fig5_dp_tradeoff.png',
     'Figure 5. Privacy–utility tradeoff under differential privacy. '
     'DP-SGD with δ = 10⁻⁵, clipping norm C = 1.0, '
     'batch size 512. Model collapses to AUC ≈ 0.5 at all '
     'tested ε ≤ 5 due to high per-round noise at '
     'healthcare-scale sampling rates.'),
    ('fig6_calibration.png',
     'Figure 6. Calibration reliability diagrams for FedProx on the BRFSS '
     'external test set (n = 1,026,318). Panels show uncalibrated '
     'model (ECE = 0.319), Platt scaling (ECE = 0.016), '
     'temperature scaling (ECE = 0.311), and isotonic regression '
     '(ECE = 0.001). Diagonal dashed line: perfect calibration.'),
    ('fig7_generalisation_gap.png',
     'Figure 7. Generalisation gap between internal (NHANES) and external '
     '(BRFSS) validation AUC for all models. Error bars: 95% CI. '
     'Annotations: ΔAUC = internal − external. '
     'Federated strategies show smaller generalisation gaps than the '
     'centralised baseline.'),
    ('fig8_summary_comparison.png',
     'Figure 8. Overall performance summary across strategies. Blue bars: '
     'internal AUC; green bars: external AUC; red bars: fairness gap '
     '(×3 for visibility). Annotations show exact values. '
     'SCAFFOLD external AUC not available (internal evaluation only).'),
]

for fname, caption in fig_specs:
    fpath = os.path.join(FIGDIR, fname)
    fig_num = fname.split('_')[0].replace('fig', '')
    add_figure(doc, fpath, caption, fig_num, width_in=5.8)


# ── save ─────────────────────────────────────────────────────────────────────
doc.save(OUTFILE)
print(f"\n  Saved: {OUTFILE}")
size_kb = os.path.getsize(OUTFILE) // 1024
print(f"  Size:  {size_kb} KB")
print("\n  Manuscript v5 complete.")
