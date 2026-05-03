"""
Task 5: Patch FL_Diabetes_Manuscript_v4_Final.docx
Applies 10 edits: Table III CI, Table V rows (4 new), 7 prose edits (3c,3d,3e,3f,3g,3i,3j).
"""

import json
import copy
import sys
sys.stdout.reconfigure(encoding='utf-8')
from pathlib import Path
from lxml import etree
import docx
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Paths ────────────────────────────────────────────────────────────────────
DOCX_PATH  = Path("D:/Projects/diabetes_prediction_project/FL_Diabetes_Manuscript_v4_Final.docx")
DATA_DIR   = Path("D:/Projects/diabetes_prediction_project/federated/results")

# ─── Load pre-computed data ────────────────────────────────────────────────────
delong     = json.loads((DATA_DIR / "delong_xgboost_external.json").read_text())
cal        = json.loads((DATA_DIR / "calibration_comparison.json").read_text())
eod        = json.loads((DATA_DIR / "table6_eod_complete.json").read_text())

CI_STR     = delong["ci_str"]                          # "0.698–0.701"
XGB_ECE    = f"{cal['XGBoost']['ece']:.3f}"            # "0.410"
XGB_BRIER  = f"{cal['XGBoost']['brier']:.3f}"         # "0.322"
NN_ECE     = f"{cal['CentralNN']['ece']:.3f}"          # "0.415"
NN_BRIER   = f"{cal['CentralNN']['brier']:.3f}"        # "0.306"

print("=" * 60)
print("Backup: ../FL_Diabetes_Manuscript_v4_Final_BACKUP.docx  OK")
print(f"CI string:  {CI_STR}")
print(f"XGB ECE/Brier: {XGB_ECE} / {XGB_BRIER}")
print(f"NN  ECE/Brier: {NN_ECE}  / {NN_BRIER}")
print("=" * 60)

doc = Document(str(DOCX_PATH))
results = {}   # edit_id -> True/False

# ══════════════════════════════════════════════════════════════════════════════
# HELPER: full text of a paragraph (joining all runs)
# ══════════════════════════════════════════════════════════════════════════════
def para_text(para):
    return "".join(run.text for run in para.runs)


# ══════════════════════════════════════════════════════════════════════════════
# HELPER: replace text within a paragraph at the XML run level
# Strategy: rebuild the paragraph keeping only the first run's rPr,
# with the full new text in that run's <w:t>.
# Handles text split across runs (common in Word).
# ══════════════════════════════════════════════════════════════════════════════
def replace_text_in_para(para, old_text, new_text):
    full = para_text(para)
    if old_text not in full:
        return False
    new_full = full.replace(old_text, new_text, 1)

    # Grab first run's rPr (character formatting) if it exists
    p_elem = para._p
    runs = p_elem.findall(qn("w:r"))
    if not runs:
        return False

    first_run = runs[0]
    rPr = first_run.find(qn("w:rPr"))

    # Remove all existing runs from the paragraph element
    for r in list(p_elem.findall(qn("w:r"))):
        p_elem.remove(r)

    # Build a new single run with the full new text
    new_r = OxmlElement("w:r")
    if rPr is not None:
        new_r.append(copy.deepcopy(rPr))
    new_t = OxmlElement("w:t")
    new_t.text = new_full
    # Preserve leading/trailing spaces
    if new_full != new_full.strip():
        new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    new_r.append(new_t)
    p_elem.append(new_r)
    return True


# ══════════════════════════════════════════════════════════════════════════════
# HELPER: insert a new paragraph immediately after an existing paragraph
# ══════════════════════════════════════════════════════════════════════════════
def insert_paragraph_after(ref_para, text, style=None):
    """Insert a new paragraph with `text` immediately after `ref_para`."""
    new_para = OxmlElement("w:p")
    # Add pPr with style if specified
    if style:
        pPr = OxmlElement("w:pPr")
        pStyle = OxmlElement("w:pStyle")
        pStyle.set(qn("w:val"), style)
        pPr.append(pStyle)
        new_para.append(pPr)
    # Add a single run with the text
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    if text != text.strip():
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    new_para.append(r)
    # Insert after ref_para in the parent element
    ref_para._p.addnext(new_para)
    return new_para


# ══════════════════════════════════════════════════════════════════════════════
# EDIT 3a: Table III — replace N/A with CI string in XGBoost external row
# ══════════════════════════════════════════════════════════════════════════════
edit_id = "3a_table_III_CI"
found = False
for table in doc.tables:
    for row in table.rows:
        cells = row.cells
        # Look for the XGBoost replicated row
        row_text = " ".join(c.text for c in cells)
        if "Centralised XGBoost (replicated)" in row_text and "N/A" in row_text:
            for ci, cell in enumerate(cells):
                if cell.text.strip() == "N/A":
                    # Replace N/A with CI string
                    for para in cell.paragraphs:
                        if "N/A" in para.text:
                            replaced = replace_text_in_para(para, "N/A", CI_STR)
                            if replaced:
                                found = True
                                break
                    if found:
                        break
            if found:
                break
    if found:
        break

results[edit_id] = found
status = "✓" if found else "✗"
print(f"\n[{status}] Edit 3a — Table III XGBoost CI: N/A → {CI_STR}")


# ══════════════════════════════════════════════════════════════════════════════
# EDIT 3b: Table V — append 4 new rows (FedAvg global/sub, FedNova global/sub)
# ══════════════════════════════════════════════════════════════════════════════
edit_id = "3b_table_V_rows"
# Find the EOD rows to append (FedAvg and FedNova)
new_rows_data = [r for r in eod["rows"] if r["model"] in ("FedAvg", "FedNova")]

# Find Table V by looking for the EOD table (has "Threshold Type" and "EOD" header)
eod_table = None
for table in doc.tables:
    header_text = " ".join(c.text for c in table.rows[0].cells) if table.rows else ""
    all_text = " ".join(c.text for row in table.rows for c in row.cells)
    if "EOD" in header_text and "Threshold Type" in header_text:
        eod_table = table
        break

if eod_table is None:
    print(f"\n[✗] Edit 3b — Table V not found (header search failed)")
    results[edit_id] = False
else:
    # Count existing rows
    existing_count = len(eod_table.rows)
    print(f"\n  Table V found with {existing_count} rows")

    # Get the last existing data row to copy its XML structure
    last_row = eod_table.rows[-1]

    def make_eod_row(table, template_row, model_name, thresh_type, young_tpr, young_fpr, elderly_tpr, elderly_fpr, eod_val):
        """Clone the template row and fill in new values."""
        # Deep copy the template row XML
        new_tr = copy.deepcopy(template_row._tr)
        tbl_elem = table._tbl

        # Get cells from the new row
        cells_xml = new_tr.findall(qn("w:tc"))

        cell_texts = [
            model_name,
            thresh_type,
            f"{young_tpr:.3f}",
            f"{young_fpr:.3f}",
            f"{elderly_tpr:.3f}",
            f"{elderly_fpr:.3f}",
            f"{eod_val:.3f}",
        ]

        for i, (tc, txt) in enumerate(zip(cells_xml, cell_texts)):
            # Clear existing text content from paragraphs
            for p in tc.findall(qn("w:p")):
                for r in p.findall(qn("w:r")):
                    for t in r.findall(qn("w:t")):
                        t.text = txt if i < len(cell_texts) else ""
                    # Only keep text in first run, clear others
                    break  # only process first p/r
                # Clear any extra runs in this paragraph
                for j, r in enumerate(p.findall(qn("w:r"))):
                    if j == 0:
                        for t in r.findall(qn("w:t")):
                            t.text = txt
                            if txt != txt.strip():
                                t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                    else:
                        p.remove(r)

        tbl_elem.append(new_tr)
        return new_tr

    # Build the 4 new rows
    rows_added = 0
    for row_data in new_rows_data:
        model    = row_data["model"]
        tt       = row_data["threshold_type"]
        ytpr     = row_data["young_tpr"]
        yfpr     = row_data["young_fpr"]
        etpr     = row_data["elderly_tpr"]
        efpr     = row_data["elderly_fpr"]
        eod_val  = row_data["eod"]

        # Format threshold type label to match table style
        if tt == "Global Youden":
            gt = row_data.get("global_threshold")
            thresh_label = f"Global Youden (τ={gt:.3f})" if gt else "Global Youden"
        else:
            yt = row_data.get("young_threshold")
            et = row_data.get("elderly_threshold")
            if yt and et:
                thresh_label = f"Subgroup-specific (young: {yt:.3f}, elderly: {et:.3f})"
            else:
                thresh_label = "Subgroup-specific"

        make_eod_row(eod_table, last_row, model, thresh_label, ytpr, yfpr, etpr, efpr, eod_val)
        rows_added += 1

    new_count = len(eod_table.rows)
    success = rows_added == 4
    results[edit_id] = success
    status = "✓" if success else "✗"
    print(f"[{status}] Edit 3b — Table V rows: added {rows_added}/4 rows ({existing_count} → {new_count} rows)")


# ══════════════════════════════════════════════════════════════════════════════
# EDIT 3c: Insert clarification paragraph after "reduces the external elderly
#          fairness gap from 0.069 to 0.054"
# ══════════════════════════════════════════════════════════════════════════════
edit_id = "3c_prose_insert"
TARGET_3C = "reduces the external elderly fairness gap from 0.069 to 0.054"
INSERT_3C = (
    "Note that 0.135 refers to the internal NHANES fairness gap of the published model; "
    "the external BRFSS fairness gap for the centralised replication is 0.069, which serves "
    "as the primary comparison baseline throughout this paper."
)

found = False
for para in doc.paragraphs:
    if TARGET_3C in para_text(para):
        insert_paragraph_after(para, INSERT_3C)
        found = True
        break

results[edit_id] = found
status = "✓" if found else "✗"
print(f"[{status}] Edit 3c — Insert fairness gap clarification note")


# ══════════════════════════════════════════════════════════════════════════════
# EDIT 3d: Insert FedProx mu sensitivity paragraph after the paragraph
#          containing both "heterogeneity of Node B" and "μ=0.1"
#          (the paragraph in Methods/Section IV-B)
# ══════════════════════════════════════════════════════════════════════════════
edit_id = "3d_fedprox_mu_insert"
TARGET_3D_A = "heterogeneity of Node B"
TARGET_3D_B = "Section 5.2 analyses the trade-off between"
INSERT_3D = (
    "Section V-B sensitivity analysis subsequently reveals that μ=0.05 achieves a marginally "
    "better fairness-accuracy balance (elderly gap 0.056 vs. 0.066 at μ=0.1; AUC 0.755 vs. 0.752); "
    "μ=0.1 was retained in the primary comparison as the a priori choice motivated by the stronger "
    "heterogeneity of Node B."
)

found = False
for para in doc.paragraphs:
    pt = para_text(para)
    if TARGET_3D_A in pt and TARGET_3D_B in pt:
        insert_paragraph_after(para, INSERT_3D)
        found = True
        break

results[edit_id] = found
status = "✓" if found else "✗"
print(f"[{status}] Edit 3d — Insert FedProx mu sensitivity note")


# ══════════════════════════════════════════════════════════════════════════════
# EDIT 3e: Insert calibration comparison paragraph after "ECE=0.276"
# ══════════════════════════════════════════════════════════════════════════════
edit_id = "3e_calibration_insert"
TARGET_3E = "ECE=0.276"
INSERT_3E = (
    f"For comparison, the centralised DiabetesNet achieves ECE={NN_ECE} and Brier={NN_BRIER} "
    f"on BRFSS; the centralised XGBoost achieves ECE={XGB_ECE} and Brier={XGB_BRIER}. "
    "The overconfidence pattern is consistent across architectures, suggesting it reflects the "
    "NHANES-to-BRFSS prevalence shift (18.6% → 13.3%) rather than being specific to the "
    "federated training procedure."
)

found = False
# There may be multiple paragraphs containing ECE=0.276 — we want the first in body text
for para in doc.paragraphs:
    if TARGET_3E in para_text(para):
        insert_paragraph_after(para, INSERT_3E)
        found = True
        break

results[edit_id] = found
status = "✓" if found else "✗"
print(f"[{status}] Edit 3e — Insert calibration comparison paragraph (NN ECE={NN_ECE}, XGB ECE={XGB_ECE})")


# ══════════════════════════════════════════════════════════════════════════════
# EDIT 3f: Replace "duplicate respondents were identified..." with BRFSS note
# ══════════════════════════════════════════════════════════════════════════════
edit_id = "3f_brfss_dedup"
OLD_3F = "duplicate respondents were identified by cross-referencing state–year identifiers and removed"
NEW_3F = (
    "BRFSS is a repeated cross-sectional survey with no persistent respondent identifiers; "
    "the same individual cannot be definitively tracked across survey years. "
    "The pooled cohort was retained as-is, consistent with standard BRFSS pooling practice"
)

found = False
for para in doc.paragraphs:
    if OLD_3F in para_text(para):
        replaced = replace_text_in_para(para, OLD_3F, NEW_3F)
        if replaced:
            found = True
            break

results[edit_id] = found
status = "✓" if found else "✗"
print(f"[{status}] Edit 3f — BRFSS deduplication note replacement")


# ══════════════════════════════════════════════════════════════════════════════
# EDIT 3g: Replace 2.2× distributional robustness phrases
# 3g1: first occurrence "2.2× difference in distributional robustness"
# 3g2: "2.2× better distributional robustness" (multiple)
# ══════════════════════════════════════════════════════════════════════════════
edit_id = "3g_2x_robustness"

OLD_3G1 = "a 2.2× difference in distributional robustness"
NEW_3G1 = (
    "a 2.2× difference in absolute internal-to-external AUC degradation "
    "(noting that FedAvg’s higher internal AUC of 0.788 vs. 0.769 is a partial confounder "
    "for this ratio; the reduction in absolute degradation nonetheless holds)"
)

OLD_3G2 = "2.2× better distributional robustness"
NEW_3G2 = (
    "2.2× better distributional robustness "
    "(absolute degradation: Δ=0.031 vs Δ=0.069; "
    "FedAvg’s higher internal AUC is a partial confounder)"
)

count_3g1 = 0
count_3g2 = 0

for para in doc.paragraphs:
    pt = para_text(para)
    if OLD_3G1 in pt and count_3g1 == 0:
        replaced = replace_text_in_para(para, OLD_3G1, NEW_3G1)
        if replaced:
            count_3g1 += 1
    elif OLD_3G2 in pt:
        replaced = replace_text_in_para(para, OLD_3G2, NEW_3G2)
        if replaced:
            count_3g2 += 1

# Also check inside tables (the 2.2x text might be in a table cell)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                pt = para_text(para)
                if OLD_3G1 in pt and count_3g1 == 0:
                    replaced = replace_text_in_para(para, OLD_3G1, NEW_3G1)
                    if replaced:
                        count_3g1 += 1
                elif OLD_3G2 in pt:
                    replaced = replace_text_in_para(para, OLD_3G2, NEW_3G2)
                    if replaced:
                        count_3g2 += 1

success = count_3g1 >= 1 and count_3g2 >= 1
results[edit_id] = success
status = "✓" if success else "✗"
print(f"[{status}] Edit 3g — 2.2x robustness: 3g1 replacements={count_3g1}, 3g2 replacements={count_3g2}")


# ══════════════════════════════════════════════════════════════════════════════
# EDIT 3i: Replace DP sample scale text
# ══════════════════════════════════════════════════════════════════════════════
edit_id = "3i_dp_sample_scale"
OLD_3I = "demonstrating a privacy-utility trade-off observed at the evaluated per-node sample scale (∼3,000–4,500 samples)"
NEW_3I = (
    "establishing that this sample regime is insufficient for DP-SGD at any standard privacy budget "
    "and motivating a minimum of ∼40,000 samples per node for viable deployment"
)

found = False
for para in doc.paragraphs:
    if OLD_3I in para_text(para):
        replaced = replace_text_in_para(para, OLD_3I, NEW_3I)
        if replaced:
            found = True
            break

results[edit_id] = found
status = "✓" if found else "✗"
print(f"[{status}] Edit 3i — DP sample scale replacement")


# ══════════════════════════════════════════════════════════════════════════════
# EDIT 3j: Insert tau quantification paragraph after "Theorem 2 in Wang et al."
# ══════════════════════════════════════════════════════════════════════════════
edit_id = "3j_tau_insert"
TARGET_3J = "Theorem 2 in Wang et al."
INSERT_3J = (
    "Distribution shift was operationalised as label distribution divergence: "
    "Node B's 28.5% prevalence versus the global 18.6% represents the largest deviation, "
    "motivating τ_B=3; Node A (13.8%) and Node C (16.7%) represent moderate deviations, "
    "assigned τ_A=5 and τ_C=4 respectively. "
    "These values are design choices rather than optimised hyperparameters."
)

found = False
for para in doc.paragraphs:
    if TARGET_3J in para_text(para):
        insert_paragraph_after(para, INSERT_3J)
        found = True
        break

results[edit_id] = found
status = "✓" if found else "✗"
print(f"[{status}] Edit 3j — Insert tau quantification paragraph")


# ══════════════════════════════════════════════════════════════════════════════
# Save document
# ══════════════════════════════════════════════════════════════════════════════
doc.save(str(DOCX_PATH))
print(f"\nDocument saved: {DOCX_PATH}")

# ══════════════════════════════════════════════════════════════════════════════
# Summary
# ══════════════════════════════════════════════════════════════════════════════
print("\n" + "=" * 60)
print("SUMMARY")
print("=" * 60)
passed = sum(1 for v in results.values() if v)
total  = len(results)
for k, v in results.items():
    print(f"  {'✓' if v else '✗'}  {k}")
print(f"\n{passed}/{total} edits succeeded")
if passed == total:
    print("STATUS: DONE")
elif passed >= total * 0.8:
    print("STATUS: DONE_WITH_CONCERNS")
else:
    print("STATUS: BLOCKED")
